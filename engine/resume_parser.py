"""Resume file parser — extracts raw text from PDF, DOCX, and TXT files."""
import io
import logging
import re

logger = logging.getLogger(__name__)


def parse_pdf(file_stream: io.BytesIO) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_stream)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        full_text = "\n".join(pages)
        if not full_text.strip():
            logger.warning("PDF parsed but no text extracted (may be scanned/image PDF)")
        return full_text
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        raise ValueError(f"Could not parse PDF: {e}")


def parse_docx(file_stream: io.BytesIO) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(file_stream)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise ValueError(f"Could not parse DOCX: {e}")


def parse_txt(file_stream: io.BytesIO) -> str:
    """Extract text from a plain text file."""
    try:
        raw = file_stream.read()
        # Try UTF-8 first, then latin-1 as fallback
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error(f"TXT parsing failed: {e}")
        raise ValueError(f"Could not parse text file: {e}")


def parse_resume(file_stream: io.BytesIO, filename: str) -> str:
    """
    Parse a resume file and return its raw text.

    Args:
        file_stream: File-like object (BytesIO)
        filename: Original filename (used to detect format)

    Returns:
        Raw text content of the resume
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    logger.info(f"Parsing resume: {filename} (format: {ext})")

    if ext == "pdf":
        text = parse_pdf(file_stream)
    elif ext in ("docx", "doc"):
        text = parse_docx(file_stream)
    elif ext == "txt":
        text = parse_txt(file_stream)
    else:
        raise ValueError(f"Unsupported file format: .{ext}")

    # Basic post-processing
    text = re.sub(r"\r\n", "\n", text)         # Normalize line endings
    text = re.sub(r"\n{3,}", "\n\n", text)     # Collapse excessive newlines
    text = text.strip()

    if len(text) < 20:
        raise ValueError(
            "Could not extract meaningful text from the resume. "
            "Please ensure the file is not empty or image-based."
        )

    # Document validation check
    is_valid, err_msg = validate_resume_content(text)
    if not is_valid:
        logger.warning(f"Document validation failed: {err_msg}")
        raise ValueError(err_msg)

    logger.info(f"Extracted {len(text)} characters from resume")
    return text


def validate_resume_content(text: str) -> tuple[bool, str]:
    """
    Intelligently checks if the uploaded document is a genuine resume or CV.
    Filters out receipts, bank statements, articles, general letters, recipes, code, and other wrong documents.
    """
    text_lower = text.lower()
    words = text_lower.split()
    word_count = len(words)
    
    # 1. Minimum length validation
    if word_count < 40:
        return False, "Wrong document uploaded: The text is too short to be a valid resume/CV."

    # 2. Check for Transactional Documents (Receipts, Bills, Invoices, bank statements)
    transactional_keywords = [
        "invoice", "receipt", "amount paid", "transaction id", "payment method",
        "payment mode", "billing address", "bill to", "tax invoice", "total paid", 
        "fees paid", "payment receipt", "receipt no", "receipt number", "order summary",
        "fee structures", "college fee", "tuition fee", "bank statement", "account balance"
    ]
    transaction_matches = sum(1 for kw in transactional_keywords if kw in text_lower)
    
    # 3. Check for Cover Letters / General Letters (which are wrong docs, not resumes)
    letter_markers = [
        "dear hiring manager", "dear recruiter", "i am writing to apply", "application for the position",
        "sincerely,", "respectfully yours", "letter of application", "to whom it may concern"
    ]
    letter_matches = sum(1 for kw in letter_markers if kw in text_lower)
    
    # 4. Check for code files (Python, HTML, JS, C++, SQL - highly technical but not resumes)
    code_markers = [
        "import ", "def ", "class ", "function()", "const ", "let ", "var ", "include <", 
        "public static void main", "select * from", "console.log"
    ]
    code_matches = sum(1 for kw in code_markers if kw in text_lower)

    # 5. Core Resume Sections (Strong signals of a real resume structure)
    resume_sections = {
        "education": ["education", "academic profile", "academic background", "study history"],
        "experience": ["experience", "work history", "employment", "professional background", "career history", "work experience"],
        "skills": ["skills", "technical skills", "competencies", "core strengths", "expertise"],
        "projects": ["projects", "personal projects", "academic projects", "key projects"],
        "contact": ["email", "phone", "contact", "linkedin", "github", "address", "portfolio"]
    }
    
    section_score = 0
    matched_sections = []
    for section_name, keywords in resume_sections.items():
        found = False
        for kw in keywords:
            if re.search(rf"\b{kw}\b", text_lower):
                found = True
                break
        if found:
            section_score += 1
            matched_sections.append(section_name)

    # 6. Career and Resume terms (Contextual signals)
    career_terms = [
        "resume", "curriculum vitae", "c.v.", "cv", "developer", "engineer", "designer", 
        "manager", "analyst", "intern", "degree", "university", "college", "school", 
        "bachelor", "master", "gpa", "cgpa", "certifications", "achievements", "objective"
    ]
    career_matches = sum(1 for term in career_terms if re.search(rf"\b{term}\b", text_lower))

    # --- DECISION HEURISTICS ---
    
    # Check 1: If it's explicitly a transactional document
    if transaction_matches >= 2 and section_score <= 1:
        return False, "Wrong document uploaded: The document appears to be a fee receipt, invoice, or financial statement."

    # Check 2: If it's a cover letter instead of a resume
    if letter_matches >= 2 and section_score <= 1:
        return False, "Wrong document uploaded: The document appears to be a cover letter. Please upload your main Resume or CV."

    # Check 3: If it's a raw source code file
    if code_matches >= 3 and section_score == 0:
        return False, "Wrong document uploaded: The document appears to be a programming source code file."

    # Check 4: General non-resume document filter
    if section_score < 2 and career_matches < 4:
        return False, (
            "Wrong document uploaded: The document does not appear to be a resume. "
            "Please ensure you are uploading a valid resume containing standard sections "
            "such as Education, Experience, and Skills."
        )

    return True, ""
